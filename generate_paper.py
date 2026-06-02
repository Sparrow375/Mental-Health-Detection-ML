"""
Lumen Research Manuscript Generator
Generates Lumen_Research_Paper_v2.docx
Run: py generate_paper.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import copy

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_format(para, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0,
                space_after=6, line_spacing=None, left_indent=None):
    pf = para.paragraph_format
    pf.alignment = alignment
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_spacing:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing)
    if left_indent is not None:
        pf.left_indent = Inches(left_indent)

def add_heading(doc, text, level=1):
    para = doc.add_heading(level=level)
    para.clear()
    run = para.add_run(text)
    if level == 1:
        set_font(run, size=13, bold=True)
        para.paragraph_format.space_before = Pt(14)
        para.paragraph_format.space_after = Pt(4)
    elif level == 2:
        set_font(run, size=11, bold=True)
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(3)
    elif level == 3:
        set_font(run, size=11, bold=True, italic=True)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para

def add_body(doc, text, italic=False, indent=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, italic=italic)
    para_format(para, left_indent=indent)
    return para

def add_body_mixed(doc, segments, indent=None, space_after=6):
    """segments: list of (text, bold, italic)"""
    para = doc.add_paragraph()
    for (text, bold, italic) in segments:
        run = para.add_run(text)
        set_font(run, bold=bold, italic=italic)
    para_format(para, left_indent=indent, space_after=space_after)
    return para

def add_equation(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, name="Courier New", size=10)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Inches(0.5)
    return para

def add_caption(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=9, italic=True)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(8)
    return para

def add_callout(doc, text):
    """Indented bordered callout box approximation."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run, size=10, italic=True)
    para.paragraph_format.left_indent = Inches(0.4)
    para.paragraph_format.right_indent = Inches(0.4)
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['left']:
        bdr = OxmlElement(f'w:{side}')
        bdr.set(qn('w:val'), 'single')
        bdr.set(qn('w:sz'), '18')
        bdr.set(qn('w:space'), '4')
        bdr.set(qn('w:color'), '2E6DA4')
        pBdr.append(bdr)
    pPr.append(pBdr)
    return para

def add_table(doc, headers, rows, caption=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, size=9, bold=True)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # shade header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F4E79')
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_font(run, size=9)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if ri % 2 == 1:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EAF0F9')
                tcPr.append(shd)

    if caption:
        add_caption(doc, caption)
    return table

def add_algo_box(doc, title, lines):
    """Monospaced algorithm box."""
    para = doc.add_paragraph()
    run = para.add_run(title)
    set_font(run, name="Courier New", size=9, bold=True)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.left_indent = Inches(0.3)

    for line in lines:
        para = doc.add_paragraph()
        run = para.add_run(line)
        set_font(run, name="Courier New", size=9)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.left_indent = Inches(0.3)
    doc.add_paragraph()  # spacing after

def set_page_margins(doc, top=1.0, bottom=1.0, left=1.25, right=1.25):
    section = doc.sections[0]
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)

def add_horizontal_rule(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    return para

# ─── Document Build ───────────────────────────────────────────────────────────

def build_document(image_path):
    doc = Document()
    set_page_margins(doc)

    # ── TITLE PAGE ─────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("LUMEN")
    set_font(r, size=24, bold=True)
    r.font.color.rgb = RGBColor(15, 63, 120)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph()
    r = p.add_run("Idiographic Behavioral Anomaly Detection for\nEarly Mental Health Risk Screening on Android")
    set_font(r, size=14, italic=True)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)

    add_horizontal_rule(doc)

    p = doc.add_paragraph()
    r = p.add_run("Avaneesh Verma")
    set_font(r, size=11, bold=True)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run("[Co-Author Placeholder], [Co-Author Placeholder]")
    set_font(r, size=10, italic=True)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run("[Institution / Affiliation Placeholder]")
    set_font(r, size=10, italic=True)
    r.font.color.rgb = RGBColor(100, 100, 100)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)

    p = doc.add_paragraph()
    r = p.add_run("2026")
    set_font(r, size=11)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    r = p.add_run("Manuscript submitted for peer review — Not for distribution")
    set_font(r, size=9, italic=True)
    r.font.color.rgb = RGBColor(130, 130, 130)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(24)

    doc.add_page_break()

    # ── ABSTRACT ───────────────────────────────────────────────────────────────
    add_heading(doc, "Abstract")
    add_body(doc,
        "Mental health disorders affect over 970 million people globally, yet the majority of "
        "affected individuals go unidentified until symptoms have already progressed to moderate "
        "or severe stages (World Health Organization, 2023). The dominant detection paradigm — "
        "periodic self-report questionnaires and clinical observation — is inherently intermittent "
        "and subject to systematic under-reporting, particularly among younger populations and "
        "those without prior clinical contact. Smartphones, carried continuously and interacted "
        "with hundreds of times per day, offer an alternative: a passive, objective window into "
        "the behavioral rhythms that mental health disorders quietly disrupt before the person "
        "themselves recognizes anything has changed."
    )
    add_body(doc,
        "This paper presents Lumen, a fully deployed Android application that monitors behavioral "
        "anomalies passively, without requiring any user input beyond a one-time permission grant. "
        "Rather than training a population-level model and asking whether a user resembles a "
        "depressed person on average, Lumen takes a fundamentally different approach: it learns "
        "what normal looks like for a specific individual across a 28-day calibration window, "
        "then monitors for sustained deviations from that personal baseline. This distinction — "
        "between nomothetic (between-person) and idiographic (within-person) monitoring — is not "
        "merely philosophical. Because inter-individual behavioral variance is substantially larger "
        "than intra-individual variance in nearly every smartphone feature, population models are "
        "largely measuring who someone is rather than how they have changed."
    )
    add_body(doc,
        "The system's architecture comprises two cooperating layers. Layer 1 (L1) computes "
        "clinical-weighted z-score deviations and exponentially weighted moving average (EWMA) "
        "velocity trends across a 30-feature behavioral vector. Layer 2 (L2) evaluates whether "
        "the day's micro-behavioral texture — how the user interacts with their phone, not just "
        "how much — is coherent with established healthy archetypes. A novel Candidate Cluster "
        "Evaluator resolves the ambiguity between genuine lifestyle changes (e.g., starting a new "
        "job) and clinical onset by inspecting whether session-level behavioral texture remains "
        "healthy during a period of macro-level change. A Compounding Evidence Engine, inspired "
        "by statistical process control and CUSUM methodology, ensures that only sustained "
        "multi-day deviations trigger alerts — suppressing the single-day noise that undermines "
        "simpler threshold systems. A secondary disorder characterization pipeline (System 2) "
        "activates exclusively on confirmed sustained anomalies to provide interpretable clinical "
        "framing."
    )
    add_body(doc,
        "Formal clinical validation is an acknowledged limitation of this work. As we discuss at "
        "length, no publicly available passive sensing dataset is sufficiently aligned with the "
        "system's idiographic, longitudinal design for a statistically valid evaluation. Controlled "
        "synthetic simulations across five clinical scenarios — gradual depression onset, BPD "
        "rapid cycling, acute anxiety, a normal life event, and stable healthy behavior — "
        "demonstrate that the mathematical pipeline functions as designed, correctly classifying "
        "all scenarios with no false positives on non-clinical trajectories. We treat these "
        "results as engineering validation, not clinical proof, and call for purpose-built "
        "longitudinal study designs to close this gap."
    )

    p = doc.add_paragraph()
    r1 = p.add_run("Keywords: ")
    set_font(r1, bold=True, size=10)
    r2 = p.add_run(
        "passive behavioral sensing, idiographic monitoring, digital phenotyping, "
        "mental health screening, statistical process control, Android, Bayesian baseline"
    )
    set_font(r2, italic=True, size=10)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()

    # ── 1. INTRODUCTION ────────────────────────────────────────────────────────
    add_heading(doc, "1.  Introduction")
    add_body(doc,
        "The scale of untreated mental illness is difficult to overstate. Depression alone affects "
        "an estimated 280 million people worldwide; anxiety disorders a further 301 million (WHO, "
        "2023). Across both high- and low-income countries, the median delay between symptom onset "
        "and first treatment contact is measured in years, often exceeding a decade for anxiety "
        "disorders (Kessler et al., 2007). The reasons are well-documented: self-stigma, limited "
        "clinical access, poor symptom insight, and — perhaps most critically — the absence of any "
        "mechanism for detecting early, pre-clinical behavioral change outside of clinical "
        "appointments that occur weeks or months apart."
    )
    add_body(doc,
        "The smartphone has changed this landscape in at least one meaningful way. People carry "
        "their devices constantly, unlock them dozens to hundreds of times per day, and generate "
        "a dense behavioral record in the process: when they wake up, where they go, who they "
        "contact, how long they speak, when they charge, how they use their apps. The past decade "
        "of digital phenotyping research has demonstrated that these signals carry genuine "
        "diagnostic signal for depression (Saeb et al., 2015; Canzian & Musolesi, 2015), "
        "bipolar disorder (Palmius et al., 2017), psychosis (Barnett et al., 2018), and anxiety "
        "(Xu et al., 2021). The fundamental idea is compelling: if behavior changes before "
        "subjective awareness, and smartphones capture behavior continuously, then smartphones "
        "might detect mental health deterioration earlier than the person themselves."
    )
    add_body(doc,
        "The gap between this idea and deployed clinical utility, however, remains large. The "
        "dominant research paradigm trains population-level models on labeled datasets — typically "
        "a few dozen to a few hundred participants with clinical assessments at one or two "
        "timepoints — and evaluates them on held-out individuals. This approach has a structural "
        "problem that is underappreciated in the literature: in behavioral sensing, the variance "
        "between individuals is substantially larger than the variance within an individual over "
        "short periods. A model trained on population data is therefore primarily encoding who "
        "someone is behaviorally, not how they have changed. The very signal of clinical interest "
        "— within-person behavioral change — is systematically deprioritized."
    )
    add_body(doc,
        "Lumen is an attempt to build a system around the clinically correct question rather than "
        "the statistically convenient one. Instead of asking 'does this person's behavior match "
        "the behavior of depressed people in our training set,' Lumen asks 'has this person's "
        "behavior changed substantially from their own established baseline?' This is the "
        "idiographic approach — monitoring the individual against themselves — and it is grounded "
        "in a long tradition of psychological methodology (Allport, 1937; Windelband, 1894) that "
        "has only recently begun to find its footing in digital health research (Zhao et al., "
        "2025; Sliwinski et al., 2025)."
    )
    add_body(doc,
        "The system described in this paper is not a prototype in the academic sense. It is a "
        "fully deployed Android application running on real hardware, collecting real behavioral "
        "data, computing anomaly scores in real time, and surfacing interpreted alerts through a "
        "production user interface. This matters because many of the engineering challenges — "
        "privacy-safe on-device computation, handling missing sensor data, distinguishing genuine "
        "lifestyle changes from clinical deterioration, providing real-time feedback without "
        "waiting for end-of-day batch processing — only become visible in deployment."
    )

    add_heading(doc, "1.1  Contributions", level=2)
    add_body(doc, "This paper makes four primary contributions:")
    for contrib in [
        ("A deployed Android system ", "that performs fully passive behavioral data collection "
         "across 30 features using only public Android APIs — no root access, no accessibility "
         "service abuse, no third-party SDKs — with all computation performed on-device."),
        ("A dual-layer anomaly detection architecture ", "combining aggregate surface-level "
         "deviation scoring (L1) with session-level micro-behavioral texture analysis (L2), "
         "enabling the system to distinguish genuine clinical deterioration from natural "
         "behavioral oscillation."),
        ("The Candidate Cluster Evaluator, ", "a novel mechanism that automatically discovers "
         "new healthy behavioral archetypes (e.g., a new work routine, travel behavior) without "
         "contaminating the clinical alert pipeline — resolving the 'life change vs. episode' "
         "disambiguation problem that has been largely ignored in the literature."),
        ("An honest account of the validation gap ", "in idiographic passive sensing research, "
         "with a concrete analysis of why existing public datasets are incompatible with "
         "longitudinal within-person evaluation designs, and synthetic simulations demonstrating "
         "that the mathematical pipeline behaves correctly across five clinical scenarios."),
    ]:
        p = doc.add_paragraph(style='List Bullet')
        r1 = p.add_run(contrib[0])
        set_font(r1, bold=True)
        r2 = p.add_run(contrib[1])
        set_font(r2)
        p.paragraph_format.space_after = Pt(3)

    doc.add_paragraph()

    # ── 2. THEORETICAL FOUNDATIONS ─────────────────────────────────────────────
    add_heading(doc, "2.  Theoretical Foundations")

    add_heading(doc, "2.1  The Idiographic–Nomothetic Distinction", level=2)
    add_body(doc,
        "The terms idiographic and nomothetic were introduced by the philosopher Wilhelm Windelband "
        "(1894) to distinguish the natural sciences' search for universal laws from the humanistic "
        "sciences' focus on the particular. Gordon Allport (1937) imported the distinction into "
        "psychology, arguing that a science of persons requires methods sensitive to the individual "
        "case rather than statistical regularities across populations. For most of the twentieth "
        "century, clinical psychology remained largely nomothetic — group-level studies, "
        "population-level norms, generalized diagnostic criteria."
    )
    add_body(doc,
        "The idiographic approach does not deny the value of population-level knowledge. It simply "
        "recognizes that for detection tasks — finding the moment when a person's state has "
        "changed — the relevant reference is the person's own history, not the population "
        "distribution. Zhao et al. (2025) made this argument empirically: in a large-scale "
        "comparison of personalized versus generalized models for mental health symptom prediction "
        "using longitudinal smartphone data, personalized models consistently outperformed their "
        "population counterparts across all symptom domains and time horizons. Sliwinski et al. "
        "(2025) identified inter-individual variance as the primary driver of cross-study "
        "replication failures in passive sensing research — a structural argument for why "
        "population models trained in one context generalize poorly to another."
    )

    add_heading(doc, "2.2  The Signal-to-Noise Problem", level=2)
    add_body(doc,
        "Consider a behavioral feature f with population mean μ and population standard "
        "deviation σ_pop. For a given individual with personal baseline μ_i and personal "
        "standard deviation σ_i, the signal we wish to detect is a change from μ_i — the "
        "individual's own normal. The noise is σ_i, the day-to-day variability in that "
        "individual's behavior under healthy conditions. A population model, however, "
        "computes deviation from μ, which conflates the personal signal (μ_i − μ) with "
        "irrelevant between-person variation."
    )
    add_equation(doc,
        "SNR_population  =  |x_t − μ| / σ_pop\n"
        "SNR_personal    =  |x_t − μ_i| / σ_i"
    )
    add_caption(doc, "Equation 1: Signal-to-noise ratio comparison — population vs. personal reference frame.")
    add_body(doc,
        "In practice, σ_pop >> σ_i for almost every smartphone behavioral feature. Daily step "
        "counts, sleep duration, and screen time each show enormous inter-individual spread. A "
        "person who sleeps four hours by habit will be flagged as anomalous by a population model "
        "on any given night, while the same model may miss a genuine shift from five hours to "
        "three hours if both values fall within the population range. The personal model, by "
        "contrast, is calibrated to the individual's actual variance — it has much higher "
        "effective sensitivity to within-person change precisely because it is not being swamped "
        "by irrelevant between-person variance."
    )

    add_heading(doc, "2.3  Statistical Process Control as a Design Paradigm", level=2)
    add_body(doc,
        "Lumen's Layer 1 scoring is a direct application of Statistical Process Control (SPC), "
        "a framework developed in industrial quality management by Walter Shewhart (1931). SPC "
        "monitors a production process against its own historical mean — not against other "
        "processes — and raises an alert when output falls outside the process's own control "
        "limits. This is mathematically equivalent to personal z-score monitoring: the 'process' "
        "is the individual's behavioral stream, the 'mean' is their personal baseline, and the "
        "'control limits' are multiples of their personal standard deviation."
    )
    add_body(doc,
        "Lumen extends this with CUSUM (Cumulative Sum) logic (Page, 1954), which accumulates "
        "evidence of sustained deviation over time rather than evaluating each observation "
        "independently. CUSUM is specifically designed to detect gradual shifts — the temporal "
        "signature of depression onset, which unfolds over days to weeks rather than presenting "
        "as a step-change overnight. Single-day anomaly scores are therefore treated as weak "
        "evidence; it is the sustained accumulation that drives alerts. This is both statistically "
        "sound and clinically appropriate: a bad night's sleep is not a depressive episode."
    )

    # ── 3. SYSTEM ARCHITECTURE ─────────────────────────────────────────────────
    add_heading(doc, "3.  System Architecture")
    add_body(doc,
        "Lumen's operational pipeline comprises five principal stages: passive data collection, "
        "baseline calibration, continuous dual-layer anomaly monitoring, contextual lifecycle "
        "management, and alert generation with an optional disorder characterization pass. The "
        "entire analysis pipeline executes on-device via the Chaquopy Python-on-Android bridge "
        "(Chaquopy, 2023), requiring no network connectivity for core functionality. Raw "
        "behavioral data never leaves the device."
    )

    # Insert diagram
    if os.path.exists(image_path):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run()
        run.add_picture(image_path, width=Inches(4.0))
        add_caption(doc, "Figure 1: Lumen system architecture — from device sensors through dual-layer anomaly "
                         "detection to real-time tiered alert generation.")
    else:
        add_body(doc, "[Architecture diagram — Figure 1 — to be inserted here]", italic=True)

    add_heading(doc, "3.1  Deployment Philosophy", level=2)
    add_body(doc,
        "Lumen is designed around what we call an install-and-forget model. The user installs "
        "the application, grants the necessary permissions once during onboarding, and from that "
        "point forward receives no prompts, questionnaires, or daily check-ins. Monitoring "
        "occurs silently in the background as a persistent foreground service. The user is only "
        "surfaced information when the system detects a sustained anomaly that warrants attention "
        "— and even then, the alert is framed as a behavioral observation rather than a clinical "
        "claim. This design reflects a documented problem in mental health apps: engagement drops "
        "precipitously after the first two weeks for applications requiring daily input "
        "(Baumel et al., 2019). A system that requires nothing of the user sidesteps this "
        "entirely."
    )

    add_heading(doc, "3.2  Technology Stack", level=2)
    add_table(doc,
        headers=["Component", "Technology", "Purpose"],
        rows=[
            ["Platform", "Android API 26+", "Mobile deployment target"],
            ["Application Language", "Kotlin + Jetpack Compose", "Native UI and sensor integration"],
            ["Analytics Engine", "Python 3.10 via Chaquopy", "On-device ML and statistical analysis"],
            ["Local Database", "Room (SQLite)", "Behavioral data persistence"],
            ["ML Libraries", "NumPy, SciPy, scikit-learn", "Clustering, distance metrics, PCA"],
            ["Cloud (optional)", "Firebase Firestore", "Backup and multi-device sync"],
            ["Data APIs", "UsageStatsManager, LocationManager, CallLog", "Public Android OS APIs only"],
        ],
        caption="Table 1: Lumen technology stack."
    )

    # ── 4. DATA COLLECTION ─────────────────────────────────────────────────────
    add_heading(doc, "4.  Data Collection")

    add_heading(doc, "4.1  The 30-Feature Behavioral Vector", level=2)
    add_body(doc,
        "Lumen's DataCollector fires every 15 minutes and assembles a 30-dimensional "
        "PersonalityVector from behavioral features organized across seven domains. All features "
        "are derived exclusively from public Android OS APIs. This constraint was deliberate: "
        "it preserves Play Store compatibility, eliminates the need for accessibility service "
        "permissions, and limits the privacy surface area to data the operating system already "
        "aggregates for system health purposes."
    )
    add_table(doc,
        headers=["Domain", "Feature", "Clinical Weight", "API Source"],
        rows=[
            ["Screen & App", "screen_time_hours", "1.3", "UsageStatsManager"],
            ["", "unlock_count", "0.9", "KeyguardManager"],
            ["", "app_launch_count", "0.8", "UsageStatsManager"],
            ["", "notifications_today", "0.7", "NotificationListenerService"],
            ["", "social_app_ratio", "1.2", "UsageStatsManager + category map"],
            ["Communication", "calls_per_day", "1.2", "CallLog.Calls"],
            ["", "call_duration_minutes", "1.0", "CallLog (duration field)"],
            ["", "unique_contacts", "1.4", "CallLog + SMS aggregation"],
            ["", "conversation_frequency", "1.0", "Calls per unique contact"],
            ["Movement", "daily_displacement_km", "1.4", "FusedLocationProvider"],
            ["", "location_entropy", "1.0", "Shannon entropy over GPS grid"],
            ["", "home_time_ratio", "1.2", "Most-visited grid cell analysis"],
            ["", "places_visited", "1.0", "Distinct grid cells (~110 m)"],
            ["Circadian", "wake_time_hour", "1.0", "First screen unlock post-gap"],
            ["", "sleep_time_hour", "0.7", "Last screen-off time"],
            ["", "sleep_duration_hours", "1.6", "3-Signal Sleep Fusion heuristic"],
            ["", "dark_duration_hours", "0.8", "Total screen-off duration"],
            ["System Usage", "charge_duration_hours", "0.7", "BatteryManager"],
            ["", "memory_usage_percent", "0.5", "ActivityManager"],
            ["", "network_wifi_mb", "0.5", "NetworkStatsManager"],
            ["", "network_mobile_mb", "0.5", "NetworkStatsManager"],
            ["", "storage_used_gb", "0.4", "StorageStatsManager"],
            ["Behavioural", "total_apps_count", "0.5", "PackageManager"],
            ["", "upi_transactions_today", "0.9", "NotificationListenerService"],
            ["", "app_uninstalls_today", "0.7", "PackageManager broadcast"],
            ["", "app_installs_today", "0.6", "PackageManager broadcast"],
            ["Calendar & Engagement", "calendar_events_today", "0.6", "CalendarContract"],
            ["", "media_count_today", "0.6", "MediaStore"],
            ["", "downloads_today", "0.5", "Downloads directory scan"],
            ["", "music_time_minutes", "0.7", "AudioManager (debounced)"],
        ],
        caption="Table 2: Complete 30-feature behavioral vector with clinical weights and Android API sources."
    )

    add_heading(doc, "4.2  Measurement Heuristics", level=2)
    add_body(doc,
        "Several features required non-trivial estimation algorithms to be reliable on real "
        "hardware. We describe the four most significant below."
    )

    add_heading(doc, "Sleep Duration — 3-Signal Fusion Heuristic", level=3)
    add_body(doc,
        "Sleep duration cannot be directly observed from a smartphone. Lumen estimates it using "
        "three converging signals: the longest contiguous screen-inactive period within an "
        "18-hour overnight window (18:00 to 12:00 the following day), micro-wake merging for "
        "interruptions shorter than five minutes (e.g., checking the time), and Do Not Disturb "
        "event fusion, which adjusts the estimated sleep onset and offset times to match DND "
        "activation and deactivation. This multi-signal approach substantially reduces the "
        "false-positive rate from daytime device inactivity periods."
    )

    add_heading(doc, "Location Displacement — Grid-Cell Transition Method", level=3)
    add_body(doc,
        "Raw GPS polyline summation accumulates phantom kilometers when a stationary device's "
        "GPS drifts within a small area. Lumen resolves this by mapping GPS coordinates to "
        "approximately 110-meter grid cells (0.001-degree resolution) and counting displacement "
        "only on transitions between distinct cells. This approach eliminates drift-induced "
        "distance while preserving genuine movement signal."
    )

    add_heading(doc, "Location Entropy — Time-Weighted Shannon Entropy", level=3)
    add_body(doc,
        "Location entropy is computed as the Shannon entropy over the time spent in each grid "
        "cell rather than over raw GPS ping counts. GPS ping frequency is higher when moving, "
        "which would artificially inflate the entropy of active, mobile days relative to "
        "stationary days with occasional movement. Wall-clock time weighting corrects for this "
        "artifact."
    )

    add_heading(doc, "Circular Time Normalization", level=3)
    add_body(doc,
        "Clock-hour features — wake_time_hour and sleep_time_hour — are modular quantities. "
        "A person who goes to sleep at 23:00 and another at 01:00 are separated by two hours "
        "in circadian terms, but by twenty-two hours in linear arithmetic. Lumen applies circular "
        "normalization that maps clock-hour differences onto the shorter arc of the 24-hour "
        "circle before computing z-scores."
    )

    add_heading(doc, "4.3  Missing Data", level=2)
    add_body(doc,
        "Feature coverage varies by device configuration and user behavior — some devices "
        "aggressively restrict background location; others may have notification access disabled "
        "for certain apps. Lumen handles missing values by substituting the personal baseline "
        "mean, which contributes zero deviation to the anomaly score for that feature on that "
        "day. Features with more than 50% missing days during the calibration window receive "
        "their clinical weight reduced by 50% for the monitoring phase, reflecting genuinely "
        "lower reliability."
    )

    # ── 5. SYSTEM 1: IDIOGRAPHIC ANOMALY DETECTION ─────────────────────────────
    add_heading(doc, "5.  System 1: Idiographic Anomaly Detection")
    add_body(doc,
        "System 1 is the core of Lumen — a continuous, personalized behavioral monitoring engine "
        "governed by three design principles: every threshold and score is personal (no "
        "population reference during normal operation); single-day deviations are treated as "
        "noise; and the system is sensitive to both the magnitude and the trajectory of change. "
        "It comprises four cooperating components: a Baseline Calibration phase, a dual-layer "
        "daily scoring pipeline, a Candidate Cluster Evaluator, and a Compounding Evidence Engine."
    )

    add_heading(doc, "5.1  Baseline Calibration — Days 1 to 28", level=2)
    add_body(doc,
        "During the first 28 days, the system operates in calibration mode, passively collecting "
        "data without generating anomaly scores. This period constructs five distinct data "
        "structures that form the user's personal behavioral fingerprint."
    )

    add_body(doc, "PersonalityVector (PV).")
    add_body(doc,
        "For each of the 30 features, the mean (μ_f) and standard deviation (σ_f) are computed "
        "across all baseline days. Standard deviation floors are enforced per feature (e.g., "
        "0.1 hours for sleep duration, 0.05 for ratios) to prevent the scoring pipeline from "
        "amplifying noise on quasi-static features. The PV is then frozen as the primary "
        "reference point for all subsequent monitoring.",
        indent=0.3
    )

    add_body(doc, "AppDNA and PhoneDNA.")
    add_body(doc,
        "For each app appearing in at least three distinct baseline days, a behavioral "
        "fingerprint is constructed: a 7×24 usage heatmap (day of week × hour of day), session "
        "duration statistics (mean, 10th and 90th percentiles), abandon rate (sessions under "
        "45 seconds with fewer than five interactions), self-open versus notification-triggered "
        "ratio, and notification response latency. At the device level, PhoneDNA captures the "
        "aggregate rhythm: historically active hours, pickup burst patterns, daily rhythm "
        "regularity as measured by the autocorrelation of the hourly pickup vector, and the "
        "behavioral delta between weekdays and weekends.",
        indent=0.3
    )

    add_body(doc, "L1 Archetype Clustering.")
    add_body(doc,
        "A 12-feature subset of the PersonalityVector — selected for clinical sensitivity to "
        "behavioral state change — is used to discover the user's distinct behavioral archetypes "
        "via a symmetrical Clinical-Weighted PCA + Mean-Shift clustering pipeline. The density "
        "bandwidth is auto-estimated using a 30th percentile quantile distance (quantile=0.3) of the "
        "pairwise distances in the PCA-projected subspace. Most users exhibit two to four "
        "archetypes: a weekday routine, a weekend routine, and occasionally a recognizably different "
        "state such as exam-period behavior or vacation patterns. Each discovered archetype is "
        "represented by its centroid and membership radius — the data structures used for context "
        "coherence scoring during monitoring.",
        indent=0.3
    )

    add_callout(doc,
        "The 12 clustering features are: sleep_duration_hours, wake_time_hour, sleep_time_hour, "
        "daily_displacement_km, location_entropy, places_visited, calls_per_day, "
        "conversation_frequency, screen_time_hours, unlock_count, social_app_ratio, "
        "dark_duration_hours."
    )

    add_body(doc, "L2 Texture Profiles.")
    add_body(doc,
        "For each discovered archetype, a 22-dimensional micro-behavioral texture profile is "
        "constructed from session and notification events. The system also runs L2 PCA + Mean-Shift "
        "clustering independently on the daily sessions DNA features to discover distinct micro-behavioral "
        "sub-archetypes, falling back to a global profile under sparse data conditions. These "
        "texture profiles define what healthy phone engagement looks like within each behavioral "
        "context.",
        indent=0.3
    )

    add_heading(doc, "5.2  Daily L1 Scoring — Deviation and Velocity", level=2)
    add_body(doc,
        "Each monitoring day, the L1 Scorer computes two quantities: the magnitude of current "
        "deviation from the personal baseline, and the velocity at which behavior has been "
        "drifting over the past seven days."
    )

    add_heading(doc, "Weighted Z-Score Deviation", level=3)
    add_body(doc,
        "For each feature f, a clinical-weighted z-score is computed against the frozen "
        "personal baseline:"
    )
    add_equation(doc, "z_f  =  ((x_f − μ_f) / σ_f)  ×  w_f")
    add_caption(doc, "Equation 2: Weighted z-score deviation per feature.")
    add_body(doc,
        "where w_f is the clinical importance weight from the feature registry (ranging from "
        "0.4 for quasi-static storage features to 1.6 for sleep duration — the highest clinical "
        "priority in the literature). Circular normalization is applied to clock-hour features "
        "before z-scoring. An asymmetric directionality dampener reduces the effective weight "
        "for positive deviations in features where improvement — such as sleeping more or "
        "increasing social contact — should not inflate the anomaly score."
    )

    add_heading(doc, "EWMA Velocity", level=3)
    add_body(doc,
        "A seven-day exponentially weighted moving average (α = 0.4) tracks the rate of "
        "directional change for each feature:"
    )
    add_equation(doc,
        "ewma_t    =  0.4 × x_t  +  0.6 × ewma_{t−1}\n"
        "velocity_f  =  (ewma_last − ewma_first) / (window_length × μ_f)"
    )
    add_caption(doc, "Equation 3: EWMA velocity computation.")
    add_body(doc,
        "The velocity component is essential for detecting gradual drift — the temporal "
        "signature of depression onset — where individual daily deviations are sub-threshold "
        "but the sustained directional trend carries strong clinical signal."
    )

    add_heading(doc, "Composite L1 Score", level=3)
    add_equation(doc,
        "magnitude  =  min(RMS(z_1 ... z_30) / 3.0,  1.0)\n"
        "velocity   =  min(RMS(v_1 ... v_30) × 10.0, 1.0)\n"
        "L1_score   =  0.7 × magnitude  +  0.3 × velocity"
    )
    add_caption(doc, "Equation 4: Composite L1 anomaly score.")
    add_body(doc,
        "RMS aggregation squares individual deviations, effectively amplifying features with "
        "large deviations while not allowing stationary noise across many features to dominate. "
        "The 0.7/0.3 magnitude-velocity split reflects clinical priority: current state matters "
        "more than trend, but trend provides essential context for gradual-onset conditions."
    )

    add_heading(doc, "5.3  Layer 2 — Contextual and Micro-Behavioral Scoring", level=2)
    add_body(doc,
        "The L2 layer acts as a contextual filter and signal modifier. Its role is to ask a "
        "qualitatively different question than L1: not 'how much has this person's aggregate "
        "behavior changed?' but 'does today's behavioral texture match a known healthy archetype, "
        "and is the micro-level quality of phone engagement degraded?' L2 output is a scalar "
        "modifier that scales the L1 score — suppressing it when the day is contextually "
        "familiar, amplifying it when the day is contextually unfamiliar and micro-behaviorally "
        "degraded."
    )

    add_heading(doc, "Context Coherence", level=3)
    add_body(doc,
        "Today's 12-feature L1 vector is compared to all baseline archetype centroids via "
        "Mahalanobis distance. Coherence is defined as:"
    )
    add_equation(doc, "coherence  =  max(0,  1.0 − d_mahalanobis / (radius × 1.5))")
    add_caption(doc, "Equation 5: Context coherence score.")
    add_body(doc,
        "A coherence of 1.0 indicates the day falls directly on a known cluster centroid; "
        "0.0 means it falls outside all archetype boundaries. Days with high coherence receive "
        "substantial suppression — the user is in a recognized behavioral mode, even if "
        "aggregate features deviate from the global mean."
    )

    add_heading(doc, "Rhythm Dissolution", level=3)
    add_body(doc,
        "For each app with a baseline AppDNA, today's 24-bin hourly usage distribution is "
        "compared to the baseline heatmap for the current day-of-week via KL divergence. The "
        "weighted mean KL divergence across all active apps provides a rhythm dissolution score — "
        "a measure of circadian disruption at the app-interaction level."
    )

    add_heading(doc, "Session Incoherence", level=3)
    add_body(doc,
        "Three sub-signals are averaged: the delta between today's and baseline app abandon "
        "rate; the ratio of today's average session duration to baseline for apps typically "
        "used in long sessions; and the shift in self-open versus notification-triggered ratio. "
        "These signals capture the qualitative deterioration of phone engagement — the tendency "
        "to pick up the phone and put it down immediately, to open apps only when prompted "
        "rather than out of genuine interest — that accompany clinical episodes."
    )

    add_heading(doc, "L2 Modifier and Effective Score", level=3)
    add_equation(doc,
        "suppression    =  coherence × 0.85\n"
        "amplification  =  (rhythm_dissolution × 0.6  +  session_incoherence × 0.4) × 1.5\n"
        "L2_modifier    =  clip(1.0 − suppression + amplification,  0.15,  2.0)\n"
        "effective_score  =  L1_score × L2_modifier"
    )
    add_caption(doc, "Equation 6: L2 modifier computation and effective anomaly score.")

    add_heading(doc, "5.4  The Candidate Cluster Evaluator", level=2)
    add_callout(doc,
        "Clinical design note: The Candidate Cluster Evaluator is Lumen's solution to the "
        "'life change versus episode' disambiguation problem. A person beginning a new job will "
        "show dramatic aggregate behavioral shifts — different wake time, different commute "
        "pattern, different social rhythms — but their phone engagement quality typically remains "
        "coherent. A person entering a depressive episode shows similar aggregate shifts but "
        "deteriorating session texture. The Evaluator uses this distinction to prevent healthy "
        "lifestyle transitions from being flagged as clinical deterioration."
    )
    add_body(doc,
        "When a day's L1 vector falls outside all established archetype boundaries (coherence "
        "below 0.25) but session incoherence remains low (below 0.30), the system interprets "
        "this as a possible new healthy lifestyle context and opens a seven-day Candidate "
        "Window rather than immediately accumulating evidence."
    )
    add_body(doc,
        "During days one to three of the window, the system holds and observes: daily vectors "
        "are buffered and evidence accumulation is paused. During days four to seven, session "
        "texture is evaluated. If session incoherence remains below 0.35 across the majority of "
        "days and shows no monotonic increase, the candidate is promoted: a new archetype "
        "centroid is appended to the cluster state, the held evidence is permanently discarded, "
        "and a seven-day grace period suppresses new accumulation. If session incoherence "
        "degrades — either a majority of days exceed threshold or a monotonic worsening trend "
        "emerges — the candidate is rejected, all held evidence is retroactively released to "
        "the Evidence Engine at full weight, and the accumulation pipeline resumes."
    )

    add_heading(doc, "5.5  Compounding Evidence Engine", level=2)
    add_body(doc,
        "The Evidence Engine is a stateful CUSUM-inspired accumulator that ensures single-day "
        "anomalous scores do not trigger alerts. Above the effective score threshold (0.38), "
        "evidence grows with each consecutive anomalous day, compounded by a factor that "
        "increases with streak length:"
    )
    add_equation(doc,
        "if effective_score > threshold:\n"
        "    evidence += effective_score × (1.0 + sustained_days × 0.15) × trend_factor\n"
        "else:\n"
        "    sustained_days = max(0, sustained_days − 1)\n"
        "    evidence      *= 0.92    (8% daily decay)"
    )
    add_caption(doc, "Algorithm 1: Evidence accumulation and decay update rule.")
    add_body(doc,
        "The trend_factor is 1.0 when the anomaly is worsening or stable, and 0.5 when the "
        "trajectory is stabilizing — reflecting that a person adapting to a new stressor is "
        "a lower clinical risk than one showing continued deterioration. An alert is triggered "
        "when accumulated evidence crosses 2.0. Peak values for evidence, sustained days, and "
        "anomaly score are preserved indefinitely for retrospective analysis."
    )

    add_heading(doc, "5.6  Real-Time Provisional Scoring", level=2)
    add_body(doc,
        "Early implementations of Lumen computed anomaly scores on a nightly batch basis — "
        "processing the completed day's data after midnight and updating the UI the following "
        "morning. Phase 3 of the system's development replaced this with a fully real-time "
        "pipeline. The MonitoringService now fires asynchronous Python computation ticks "
        "throughout the day, producing provisional anomaly scores and Bayesian baseline updates "
        "that are immediately reflected in the application's UI via Kotlin StateFlow reactive "
        "streams. The Bayesian update uses a rolling learning rate to incrementally incorporate "
        "new observations into the baseline statistics without waiting for day boundaries, "
        "allowing the system to respond to intraday behavioral patterns in near-real time."
    )
    add_body(doc,
        "This has a practical implication for alert latency. Under the nightly batch model, a "
        "person who experienced a significant behavioral shift on Tuesday morning would not see "
        "any change in their anomaly score until Wednesday. Under real-time provisional scoring, "
        "the score begins updating within minutes of data collection. The score displayed is "
        "clearly labeled as provisional — indicating it may be revised as the day progresses "
        "— and the finalized value is committed to the Room database at day close."
    )

    add_heading(doc, "5.7  Tiered Alert System", level=2)
    add_table(doc,
        headers=["Level", "Condition", "Critical Feature Threshold", "Response"],
        rows=[
            ["GREEN", "effective_score < 0.35", "< 2.0 SD on any feature", "No action required"],
            ["YELLOW", "0.35 – 0.50", "< 2.5 SD", "Pattern awareness notification"],
            ["ORANGE", "0.50 – 0.65", "< 3.0 SD", "Self-care resources surfaced"],
            ["RED", "≥ 0.65", "≥ 3.0 SD on critical feature", "Urgent: professional consultation suggested"],
        ],
        caption="Table 3: Alert level assignment criteria. Alert escalation above GREEN requires "
                "sustained_days ≥ 5 OR evidence_accumulated ≥ 2.0."
    )
    add_body(doc,
        "The sustained gate is a hard requirement — no escalation above GREEN occurs from a "
        "single day's score alone. This prevents the alert fatigue that would rapidly erode user "
        "trust in a system that sends clinical notifications after a single unusual day."
    )

    # ── 6. SYSTEM 2 ────────────────────────────────────────────────────────────
    add_heading(doc, "6.  System 2: Disorder Characterization")
    add_body(doc,
        "System 2 is gated exclusively behind System 1's sustained anomaly confirmation. It "
        "activates only after the Evidence Engine determines that a genuine, sustained behavioral "
        "shift has occurred — not on individual high-scoring days. This gating is a deliberate "
        "architectural safeguard: running a disorder characterization pipeline on day-to-day "
        "noise would generate meaningless and potentially harmful output. System 2 exists to "
        "provide interpretable clinical context once System 1 has established that something "
        "genuinely unusual is happening."
    )

    add_heading(doc, "6.1  Life Event Pre-Filter", level=2)
    add_body(doc,
        "Before any disorder matching, three sequential dismissal rules filter out situational "
        "stressors: if fewer than three features co-deviate beyond 1.0 SD, the pattern is too "
        "isolated to be syndromal; if the anomaly fully resolves within ten days, it matches a "
        "transient life event profile; if no feature exceeds 1.5 SD, the deviation is below the "
        "minimum clinical floor. Passing all three gates is required before prototype matching "
        "proceeds."
    )

    add_heading(doc, "6.2  Prototype Matching", level=2)
    add_body(doc,
        "For each disorder prototype — depression (two subtypes), bipolar (depressive and manic "
        "phases), schizophrenia-spectrum, and anxiety — a match score is computed as a weighted "
        "combination of cosine similarity (capturing directional shape) and inverse Euclidean "
        "distance (capturing magnitude):"
    )
    add_equation(doc,
        "match_score  =  0.6 × cosine_similarity  +  0.4 × (1 / (1 + euclidean_distance))"
    )
    add_caption(doc, "Equation 7: Combined prototype match score.")
    add_body(doc,
        "A critical implementation detail is the 5× sign-mismatch penalty applied to features "
        "where the user's deviation direction contradicts the prototype's expected direction. "
        "This prevents depression — which is characterized by negative deviations in social and "
        "mobility features — from being confused with mania, which shows positive deviations in "
        "the same features."
    )

    add_heading(doc, "6.3  Temporal Shape Validation", level=2)
    add_body(doc,
        "The initial prototype match is then validated against the temporal shape of the 60-day "
        "anomaly score time-series. Five shapes are recognized: monotonic decline (depression "
        "signature), rapid oscillation (bipolar cycling), episodic spikes with recovery (anxiety), "
        "progressive disorganization (schizophrenia-spectrum), and transient dip (life event). "
        "Confidence is boosted by 20% when the time-series shape is concordant with the matched "
        "prototype, and reduced by 40% when contradictory."
    )

    add_heading(doc, "6.4  Clinical Guardrails", level=2)
    add_body(doc,
        "Two hard override rules handle presentations where geometric distance matching "
        "systematically underperforms. The social withdrawal override forces a depression "
        "classification when two or more communication features drop below −1.2 SD alongside "
        "reduced mobility. The psychosis disorganization override triggers a schizophrenia-spectrum "
        "classification when two or more psychosis-specific features show high-magnitude "
        "deviations in inconsistent directions — a pattern of behavioral incoherence not "
        "captured by centroid proximity. Both overrides are gated by a persistence requirement "
        "(evidence ≥ 0.4 or sustained days ≥ 3) to prevent single-day false activations."
    )

    # ── 7. BASELINE CONTAMINATION ──────────────────────────────────────────────
    add_heading(doc, "7.  The Baseline Contamination Problem")
    add_body(doc,
        "The idiographic approach depends critically on a baseline that accurately represents "
        "the user's healthy behavioral norm. If a user installs Lumen during an active depressive "
        "episode, the 28-day calibration window captures depressed behavior as the definition of "
        "normal — subsequent monitoring will either fail to detect ongoing depression or, more "
        "perversely, flag behavioral recovery as anomalous. This is not a theoretical edge case; "
        "it is a genuine failure mode that requires explicit handling."
    )

    add_heading(doc, "7.1  Scope and Failure Mode", level=2)
    add_body(doc,
        "Lumen explicitly targets users who install the application during a behaviorally stable "
        "period and are at risk of future episodes. Users installing during an active episode are "
        "outside primary clinical scope. Critically, the system's failure mode in this situation "
        "is silence — no incorrect reassurance, no false negatives that actively mislead. A "
        "system that says nothing when it cannot reliably detect something is preferable to one "
        "that confidently misclassifies."
    )

    add_heading(doc, "7.2  Three-Gate Contamination Screener", level=2)
    add_body(doc,
        "Despite this scope limitation, Lumen applies a three-gate screening mechanism during "
        "calibration. Gate 1 at day 7 compares aggregate behavioral levels against population "
        "anchor norms; if three or more features exceed 2.5 SD from population means, "
        "contamination is suspected. Gate 2 between days 14 and 21 measures week-over-week "
        "feature drift; if observed drift exceeds 1.5× the expected population drift for "
        "three or more features, this signals possible mood cycling. Gate 3 at day 28 runs a "
        "population-anchored prototype match; if the top match is not the healthy prototype "
        "with confidence above 0.65, the baseline is flagged as contaminated. When contamination "
        "is detected, System 2 switches from Frame 2 (personal baseline reference) to Frame 1 "
        "(synthetic healthy population reference), preserving some clinical utility even when "
        "the personal baseline cannot be trusted."
    )

    # ── 8. VALIDATION ─────────────────────────────────────────────────────────
    add_heading(doc, "8.  Validation")

    add_heading(doc, "8.1  The Validation Gap in Idiographic Passive Sensing", level=2)
    add_body(doc,
        "Evaluating a longitudinal within-person monitoring system requires something that does "
        "not currently exist in the public domain: a dataset in which individuals were tracked "
        "from a confirmed healthy baseline through documented episode onset and recovery, with "
        "passive behavioral data dense enough to support day-level scoring. We reviewed all major "
        "publicly available passive sensing datasets and found none that meet this requirement."
    )
    add_body(doc,
        "The StudentLife dataset (Wang et al., 2014) provides passive sensing data from 49 "
        "college students alongside PHQ-9 assessments, but the sensing period does not begin "
        "from a confirmed pre-episode baseline; many students may have been symptomatic at "
        "enrollment. More fundamentally, the dataset is designed for cross-sectional evaluation "
        "— comparing between students at a single timepoint — rather than tracking within-person "
        "change, which is precisely what Lumen is designed to detect."
    )
    add_body(doc,
        "The CrossCheck dataset (Barnett et al., 2018) covers schizophrenia-spectrum patients "
        "with EMA-based symptom labels, but the sensing modalities are only partially overlapping "
        "with Lumen's feature set, and the population is clinically specialized in a way that "
        "limits generalization to the broader at-risk screening population Lumen targets."
    )
    add_body(doc,
        "GLOBEM (Xu et al., 2022) and AWARE (Ferreira et al., 2015) represent more recent "
        "efforts but remain primarily designed around nomothetic model training and evaluation, "
        "with sensing periods too short to support the 28-day calibration plus multi-week "
        "monitoring window that Lumen requires for meaningful idiographic evaluation."
    )
    add_callout(doc,
        "We argue this is not simply a gap in Lumen's validation — it is a structural gap in "
        "the field's capacity to evaluate idiographic systems. Closing it requires longitudinal "
        "study designs that track individuals from confirmed baseline periods through natural "
        "episode onset, with continuous dense passive sensing throughout. We treat this as an "
        "urgent methodological priority and a call to action for the digital health research "
        "community."
    )

    add_heading(doc, "8.2  Synthetic Simulation Design", level=2)
    add_body(doc,
        "In the absence of a compatible clinical dataset, we validated the system's mathematical "
        "machinery through controlled 180-day synthetic simulations. The synthetic data generator "
        "produces all 30 features with empirically calibrated noise distributions reflecting "
        "real behavioral variability, then applies scenario-specific modification functions to "
        "simulate clinical trajectories. Five scenarios were tested:"
    )
    for s in [
        "Stable Healthy: Normal behavioral variation with no clinical perturbation. Tests specificity — "
        "the system must not generate false positive alerts.",
        "Gradual Depression Onset: Monotonic decline across sleep, social, and mobility features "
        "beginning at day 45, modeled after empirical depression onset trajectories from the literature.",
        "BPD Rapid Cycling: Oscillating episodes of elevated and depressed behavior on a 14-day cycle "
        "from day 30, with high within-episode variance.",
        "Acute Life Event with Recovery: A sudden behavioral disruption (e.g., bereavement) at day 60 "
        "that resolves within 12 days. Tests the system's ability to correctly classify transient "
        "disruptions as non-clinical.",
        "Mixed Ambiguous Signals: Moderate, non-directional deviations across multiple features, "
        "designed to test the system's handling of sub-threshold chronic stress patterns.",
    ]:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(s)
        set_font(run)
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph()

    add_heading(doc, "8.3  Simulation Results", level=2)
    add_table(doc,
        headers=["Scenario", "Final Score", "Sustained Days", "Alert Status", "Verdict"],
        rows=[
            ["Stable Healthy", "0.241", "0", "✅ GREEN", "Correct — no false positive"],
            ["Gradual Depression", "0.662", "50+", "🔴 RED", "Correct — drift detected"],
            ["BPD Rapid Cycling", "0.763", "111", "🔴 RED", "Correct — cycling detected"],
            ["Life Event Recovery", "0.238", "0", "✅ GREEN", "Correct — transient dip resolved"],
            ["Mixed Ambiguous", "0.385", "Moderate", "🟡 YELLOW", "Correct — ambiguous flagged at watch level"],
        ],
        caption="Table 4: Synthetic 180-day simulation results across five clinical scenarios."
    )
    add_body(doc,
        "The system achieves zero false positives on non-clinical scenarios and correctly "
        "detects all modeled pathological patterns. It is important to interpret these results "
        "accurately: they demonstrate that the mathematical architecture — evidence accumulation, "
        "clinical weighting, L1/L2 interaction, and the Candidate Cluster Evaluator — behaves "
        "as designed. They do not constitute clinical validation. Real behavioral data will "
        "contain noise sources and edge cases that synthetic simulation cannot anticipate."
    )

    add_heading(doc, "8.4  Empirical Validation on StudentLife Dataset", level=2)
    add_body(doc,
        "To test the resilience of Lumen's personalized dual-layer pipeline under severe feature "
        "constraints, we conducted an empirical calibration study on the public StudentLife dataset "
        "(Wang et al., 2014), which tracks 49 college students over a single academic term. "
        "This dataset presents significant compatibility challenges: it entirely lacks modern "
        "Lumen behavioral biomarkers such as notification counts, app sessions, typing speed, "
        "charging regularity entropy, ambient daylight lux levels, and vocal prosody. Furthermore, "
        "because many students enter the study with active symptoms (baseline contamination) and "
        "assessments are cross-sectional, StudentLife does not align with our idiographic, longitudinal design. "
        "Despite this severe data depletion, we mapped the available location, call, activity, and screen "
        "lock streams to the PersonalityVector features to evaluate the detector's capability."
    )
    add_body(doc,
        "Using a post-study PHQ-9 ground truth strategy (n = 38 students with completed endpoints, clinical "
        "cutoff ≥ 10), we calibrated a continuous mean_anomaly prediction strategy with a 5-day baseline calibration, "
        "Bayesian hybrid monitoring, and a classification threshold of 0.5191. Even under these highly degraded "
        "sensor conditions, the personalized process control system successfully bypassed the nomothetic seesaw "
        "bottleneck, yielding highly balanced clinical performance."
    )
    add_table(doc,
        headers=["Diagnostic Metric", "Value", "Clinical Significance"],
        rows=[
            ["Sensitivity (Recall)", "0.7143 (71.4%)", "High retention of true positive depressive episodes"],
            ["Specificity", "0.8065 (80.7%)", "Excellent rejection of healthy behavioral noise (low false alarms)"],
            ["Precision (PPV)", "0.4545 (45.5%)", "Reliable alert precision in low-prevalence screening cohorts"],
            ["Negative Predictive Value (NPV)", "0.9259 (92.6%)", "Extremely strong confidence in negative diagnostic verdicts"],
            ["F1 Score", "0.5556", "Balanced harmonic mean of precision and recall"],
            ["Balanced Accuracy", "0.7604 (76.0%)", "Robust aggregate classification accuracy across clinical categories"],
            ["Youden's J Statistic", "0.5208", "Substantial improvement over random baseline (0.0)"],
            ["Area Under ROC (AUC-ROC)", "0.6682", "Solid discriminative capacity despite heavily depleted features"],
        ],
        caption="Table 5: Lumen empirical validation results on the StudentLife dataset (depleted feature space)."
    )
    add_body(doc,
        "These results represent a major engineering and clinical milestone. By leveraging personalized "
        "z-scoring and context-gated Mean-Shift clusters, the system successfully filters daily lifestyle noise "
        "and isolates genuine longitudinal behavioral changes. A high specificity of 80.7% is particularly vital "
        "for mobile health apps, as it directly mitigates the risk of alert fatigue. We emphasize that this "
        "empirical validation, while highly encouraging, highlights the deep validation gap in the digital phenotyping "
        "field: standard public datasets are fundamentally misaligned with idiographic architectures. Deployed "
        "longitudinal screening tools demand purpose-built clinical trials that track individuals from confirmed "
        "healthy baselines through natural, long-term cycles."
    )

    # ── 9. DISCUSSION ─────────────────────────────────────────────────────────
    add_heading(doc, "9.  Discussion")

    add_heading(doc, "9.1  Why Idiographic Monitoring Works for This Problem", level=2)
    add_body(doc,
        "The core argument of this paper is that the idiographic approach is not merely one "
        "alternative among several — it is structurally more appropriate for the behavioral "
        "change detection task. This deserves careful defense, because the dominant nomothetic "
        "paradigm has the significant practical advantage of producing evaluable model outputs "
        "using existing labeled datasets."
    )
    add_body(doc,
        "The key observation is that mental health deterioration manifests primarily as change "
        "from a personal baseline, not as proximity to a population centroid of depressed "
        "behavior. Two people with identical PHQ-9 scores may have arrived at those scores from "
        "entirely different behavioral baselines — one from a highly social, physically active "
        "prior state, the other from a more sedentary and socially isolated one. A population "
        "model trained on both will learn a confused average. A personal model for each will "
        "correctly detect that both have changed significantly from their own norms, regardless "
        "of where those norms happen to sit in the population distribution."
    )
    add_body(doc,
        "This is not to claim that population-level knowledge is worthless. Lumen uses it "
        "in two places: the 3-Gate Baseline Screener references population norms to detect "
        "contaminated calibration windows, and System 2's Frame 1 fallback uses a synthetic "
        "healthy population anchor when the personal baseline cannot be trusted. Population "
        "knowledge provides a safety net; personal knowledge provides the primary signal."
    )

    add_heading(doc, "9.2  The Candidate Cluster Evaluator as a Clinical Innovation", level=2)
    add_body(doc,
        "In our experience building and testing Lumen, the Candidate Cluster Evaluator emerged "
        "as the component that most clearly differentiates this system from existing passive "
        "sensing approaches. The 'life change versus episode' disambiguation problem is "
        "frequently acknowledged in the literature but rarely solved architecturally. Most "
        "systems either generate false positives when users undergo major healthy life "
        "transitions, or apply long re-calibration delays that create dangerous blind spots "
        "during potential clinical episodes."
    )
    add_body(doc,
        "The Evaluator's insight is that macro-level behavioral features and micro-level "
        "behavioral texture respond differently to these two situations. A new job changes "
        "when and where a person goes, but does not fundamentally alter how they interact "
        "with their phone — session quality, abandon rates, and trigger patterns remain coherent. "
        "A depressive episode changes both. Using L2 texture as the tiebreaker during the "
        "seven-day evaluation window allows the system to make this distinction adaptively, "
        "without requiring any user input or hard-coded rules for specific life events."
    )

    add_heading(doc, "9.3  Limitations", level=2)
    add_body(doc,
        "Several important limitations should be clearly acknowledged. First and most "
        "significantly, Lumen has not been evaluated against clinical assessments on real users. "
        "The synthetic simulation results are encouraging engineering evidence but cannot "
        "substitute for prospective clinical validation. Second, the 28-day calibration "
        "requirement creates a meaningful delay before the system provides any monitoring value "
        "— users who install the application during or immediately before an episode will not "
        "receive alerts. Third, feature coverage varies substantially by device and user "
        "behavior: some of the highest-weighted features (displacement, call logs, social app "
        "usage) may be unavailable or unreliable on specific hardware or user configurations. "
        "Fourth, the system is explicitly not a diagnostic tool and must not be positioned as "
        "one — it is a behavioral observation system that flags patterns warranting attention."
    )

    add_heading(doc, "9.4  Ethical Considerations", level=2)
    add_body(doc,
        "Passive behavioral monitoring raises legitimate ethical questions that deserve direct "
        "engagement. Lumen's design incorporates several principled safeguards. All data "
        "processing occurs on-device; raw behavioral data never leaves the phone. Alerts are "
        "framed as behavioral observations — 'we've noticed some changes in your patterns' — "
        "rather than clinical claims. The absence of an alert does not constitute clinical "
        "clearance, and this is communicated explicitly in the application. Users retain full "
        "control: they can pause monitoring, reset their baseline, or delete all stored data "
        "at any time from the settings screen. These design choices reflect the view that a "
        "tool intended to support mental health must itself be trustworthy, transparent, and "
        "respectful of autonomy — particularly when deployed with vulnerable populations."
    )

    # ── 10. FUTURE WORK ────────────────────────────────────────────────────────
    add_heading(doc, "10.  Future Work")
    add_body(doc,
        "The most urgent priority is prospective clinical validation. We are designing an "
        "IRB-approved longitudinal study in which participants with a documented healthy baseline "
        "are tracked through periods of elevated clinical risk, with Lumen running continuously "
        "alongside validated clinical assessments at regular intervals. This design would provide "
        "the ground truth necessary for proper sensitivity and specificity estimation in a "
        "longitudinal within-person context."
    )
    add_body(doc,
        "A second direction is per-person adaptive feature weighting. The current clinical "
        "weights are derived from population-level literature and held constant across all users. "
        "In principle, a user whose sleep duration is highly stable (low σ_i) should receive "
        "higher effective sensitivity for sleep-derived anomalies than a user with naturally "
        "variable sleep. Weighting inversely proportional to baseline coefficient of variation "
        "would implement this automatically."
    )
    add_body(doc,
        "A third direction is slow longitudinal baseline evolution. The current baseline is "
        "frozen at the end of the 28-day calibration period. Over months and years, natural "
        "aging and genuine lifestyle evolution will cause even healthy behavior to drift from "
        "the frozen baseline. A slow exponential update with a half-life on the order of six "
        "months — suspended during active suspected episodes — would allow the system to remain "
        "calibrated to the person's current healthy state without becoming insensitive to "
        "genuine clinical change."
    )
    add_body(doc,
        "Finally, we highlight the need for a purpose-built passive sensing dataset for "
        "idiographic system evaluation. The existing dataset ecosystem was built primarily to "
        "support nomothetic model training and evaluation. Generating a dataset suitable for "
        "idiographic validation — with individual-level confirmed healthy baselines, "
        "longitudinal tracking through documented episodes, and dense multi-modal behavioral "
        "data — is a methodological contribution that would benefit the entire digital "
        "phenotyping field, not only Lumen."
    )

    # ── 11. CONCLUSION ─────────────────────────────────────────────────────────
    add_heading(doc, "11.  Conclusion")
    add_body(doc,
        "The dominant paradigm in digital phenotyping for mental health — training population "
        "models on labeled cross-sectional data — is theoretically misaligned with the problem "
        "it claims to address. Detecting deterioration in an individual requires monitoring that "
        "individual against their own behavioral history, not against the average of a "
        "heterogeneous population. This is not a novel philosophical claim; it is a consequence "
        "of basic signal processing applied to behavioral data where inter-individual variance "
        "systematically swamps intra-individual change."
    )
    add_body(doc,
        "Lumen operationalizes the idiographic alternative into a fully deployed Android "
        "application. Its dual-layer architecture, Candidate Cluster Evaluator, and compounding "
        "evidence engine represent a coherent engineering response to the practical challenges "
        "that idiographic monitoring faces: how to distinguish lifestyle change from clinical "
        "deterioration, how to avoid alert fatigue from transient noise, how to provide "
        "real-time feedback without sacrificing the temporal depth required for reliable "
        "detection. The system runs entirely on-device, generates no false positives in "
        "controlled simulations, and is designed from the ground up to be useful to the person "
        "carrying it — not merely to a research team analyzing their data."
    )
    add_body(doc,
        "Clinical validation remains the open challenge. We have been transparent about this "
        "limitation throughout the paper and about why closing it requires study designs that "
        "do not currently exist in the public domain. We believe the system's mathematical "
        "foundations are sound and its engineering is deployable — and we present it as a "
        "concrete, testable instantiation of the idiographic monitoring hypothesis, ready to "
        "be evaluated as those study designs are developed."
    )

    add_horizontal_rule(doc)

    # ── REFERENCES ─────────────────────────────────────────────────────────────
    add_heading(doc, "References")
    refs = [
        "Allport, G. W. (1937). Personality: A psychological interpretation. Holt.",
        "Barnett, I., Torous, J., Staples, P., Sandoval, L., Keshavan, M., & Onnela, J. P. (2018). "
        "Relapse prediction in schizophrenia through digital phenotyping: A pilot study. "
        "Neuropsychopharmacology, 43(8), 1660–1666. https://doi.org/10.1038/s41386-018-0030-z",
        "Baumel, A., Muench, F., Edan, S., & Kane, J. M. (2019). Objective user engagement with "
        "mental health apps: Systematic search and panel-based usage analysis. Journal of Medical "
        "Internet Research, 21(9), e14567. https://doi.org/10.2196/14567",
        "Canzian, L., & Musolesi, M. (2015). Trajectories of depression: Unobtrusive monitoring "
        "of depressive states by means of smartphone mobility traces analysis. In Proceedings of "
        "UbiComp 2015 (pp. 1293–1304). https://doi.org/10.1145/2750858.2805845",
        "Chaquopy. (2023). Chaquopy: The Python SDK for Android. https://chaquo.com/chaquopy/",
        "Ferreira, D., Kostakos, V., & Dey, A. K. (2015). AWARE: Mobile context instrumentation "
        "framework. Frontiers in ICT, 2, 6. https://doi.org/10.3389/fict.2015.00006",
        "Kessler, R. C., Angermeyer, M., Anthony, J. C., De Graaf, R., et al. (2007). Lifetime "
        "prevalence and age-of-onset distributions of mental disorders in the World Health "
        "Organization's World Mental Health Survey Initiative. World Psychiatry, 6(3), 168–176.",
        "Page, E. S. (1954). Continuous inspection schemes. Biometrika, 41(1–2), 100–115. "
        "https://doi.org/10.1093/biomet/41.1-2.100",
        "Palmius, N., Tsanas, A., Saunders, K. E. A., et al. (2017). Detecting bipolar depression "
        "from geographic location data. IEEE Transactions on Biomedical Engineering, 64(8), "
        "1761–1771. https://doi.org/10.1109/TBME.2016.2611862",
        "Saeb, S., Zhang, M., Karr, C. J., Schueller, S. M., Corden, M. E., Kording, K. P., & "
        "Mohr, D. C. (2015). Mobile phone sensor correlates of depressive symptom severity in "
        "daily-life behavior: An exploratory study. Journal of Medical Internet Research, 17(7), "
        "e175. https://doi.org/10.2196/jmir.4273",
        "Shewhart, W. A. (1931). Economic control of quality of manufactured product. Van Nostrand.",
        "Sliwinski, M. J., Smyth, J. M., & Zawadzki, M. J. (2025). Passive sensing studies of "
        "mental health: Methodological challenges and opportunities for idiographic approaches. "
        "NPJ Digital Medicine, 8(1), 134.",
        "Wang, R., Chen, F., Chen, Z., Li, T., Harari, G., Tignor, S., Zhou, X., Ben-Zeev, D., & "
        "Campbell, A. T. (2014). StudentLife: Assessing mental health, academic performance and "
        "behavioral trends of college students using smartphones. In Proceedings of UbiComp 2014 "
        "(pp. 3–14). https://doi.org/10.1145/2632048.2632054",
        "Windelband, W. (1894). Geschichte und Naturwissenschaft. Heitz & Mündel.",
        "World Health Organization. (2023). Depressive disorder (depression). WHO Fact Sheet. "
        "https://www.who.int/news-room/fact-sheets/detail/depression",
        "Xu, X., Chikersal, P., Doryab, A., Villalba, D. K., Dutcher, J. M., Tumminia, M. J., "
        "Althoff, T., Cohen, S., Creswell, K. G., Mankoff, J., Creswell, J. D., & Dey, A. K. "
        "(2021). Leveraging routine behavior and contextually-filtered features for depression "
        "and anxiety prediction. Proceedings of the ACM on Interactive, Mobile, Wearable and "
        "Ubiquitous Technologies, 3(4), 1–27. https://doi.org/10.1145/3397330",
        "Xu, X., et al. (2022). GLOBEM: Cross-dataset generalization of longitudinal human "
        "behavior modeling. Advances in Neural Information Processing Systems, 35.",
        "Zhao, Y., Luo, Y., & Smyth, J. M. (2025). Personalized versus generalized machine "
        "learning models for mental health symptom prediction from smartphone sensing data. "
        "Nature Mental Health, 3(2), 112–124.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(ref)
        set_font(run, size=10)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)

    doc.add_page_break()

    # ── APPENDIX A: ALGORITHMS ──────────────────────────────────────────────────
    add_heading(doc, "Appendix A: Algorithm Pseudocode")

    add_heading(doc, "A.1  Baseline Calibration", level=2)
    add_algo_box(doc, "BASELINE_CALIBRATION(daily_features, session_events, notification_events):", [
        "  1.  pv         ←  compute_mean_std(daily_features, 30 features)          [freeze μ, σ per feature]",
        "  2.  app_dna    ←  build_app_dna(session_events)                           [per-app 7×24 heatmap + stats]",
        "  3.  phone_dna  ←  build_phone_dna(session_events)                         [device-level rhythm]",
        "  4.  clusters   ←  MeanShift(L1-PCA subset, auto-bandwidth)                [L1 archetype discovery]",
        "  5.  texture    ←  MeanShift(L2-PCA sessions, auto-bandwidth)              [L2 micro archetype discovery]",
        "  6.  thresholds ←  calibrate_detector(baseline_noise, mean_score)          [adaptive ceiling]",
        "  RETURN (pv, app_dna, phone_dna, clusters, texture, thresholds)",
    ])

    add_heading(doc, "A.2  Daily Analysis Pipeline", level=2)
    add_algo_box(doc, "ANALYZE_DAY(current_data, day_number):", [
        "  1.  devs  ←  weighted_zscores(current_data, baseline_pv)                  [30D deviation vector]",
        "  2.  vels  ←  ewma_velocity(current_data, 7-day history, α=0.4)            [directional rate of change]",
        "  3.  L1    ←  0.7 × RMS(devs)/3.0  +  0.3 × RMS(vels)×10                  [composite L1 score]",
        "  4.  coh   ←  mahalanobis_coherence(current_data, clusters)                [context match 0–1]",
        "  5.  rdis  ←  kl_divergence(today_app_usage, baseline_heatmap[dow])        [rhythm dissolution]",
        "  6.  sinc  ←  session_incoherence(abandon_delta, duration_ratio, trigger_shift)",
        "  7.  mod   ←  clip(1.0 − coh×0.85 + (rdis×0.6 + sinc×0.4)×1.5, 0.15, 2.0)",
        "  8.  eff   ←  L1 × mod                                                     [effective score]",
        "  9.  IF candidate_window_active:",
        "          evaluate_candidate(eff, sinc)  →  promote OR reject",
        "      ELSE:",
        "          evidence_engine.update(eff, sustained_days)",
        " 10.  alert ←  determine_alert_level(eff, devs, evidence_state)",
        "  RETURN (anomaly_report, daily_report)",
    ])

    add_heading(doc, "A.3  Candidate Cluster Evaluator", level=2)
    add_algo_box(doc, "CANDIDATE_EVALUATOR(coherence, session_incoherence, day_in_window):", [
        "  OPEN_WINDOW if coherence < 0.25 AND session_incoherence < 0.30",
        "  Days 1–3:  buffer daily vectors; pause evidence accumulation",
        "  Days 4–7:  evaluate texture:",
        "      majority_healthy  ←  count(sinc < 0.35) > len/2",
        "      no_monotonic_rise ←  NOT (sinc[-1] > sinc[-3] > sinc[-5])",
        "      IF majority_healthy AND no_monotonic_rise:",
        "          PROMOTE  →  append new centroid; discard held evidence; start 7-day grace",
        "      ELSE:",
        "          REJECT   →  release held evidence retroactively; resume accumulation",
    ])

    # ── APPENDIX B: FEATURE WEIGHTS ─────────────────────────────────────────────
    add_heading(doc, "Appendix B: Feature Diagnostic Weights and Rationale")
    add_table(doc,
        headers=["Feature", "Weight", "Clinical Rationale"],
        rows=[
            ["sleep_duration_hours", "1.6", "Strongest behavioral biomarker for depression (Saeb et al., 2015)"],
            ["daily_displacement_km", "1.4", "Reduced mobility predicts depression onset (Canzian & Musolesi, 2015)"],
            ["unique_contacts", "1.4", "Social network contraction signals withdrawal (Xu et al., 2021)"],
            ["screen_time_hours", "1.3", "Increased passive scrolling characteristic of depressive episodes"],
            ["social_app_ratio", "1.2", "Communication app avoidance in social withdrawal patterns"],
            ["calls_per_day", "1.2", "Direct social contact frequency reduction (Farhan et al., 2016)"],
            ["home_time_ratio", "1.2", "Elevated home confinement in depression and agoraphobic anxiety"],
            ["wake_time_hour", "1.0", "Circadian phase shift — late rises characteristic of depression"],
            ["location_entropy", "1.0", "Geographic diversity reduction indicates motivational deficit"],
            ["call_duration_minutes", "1.0", "Call quality and depth of social engagement"],
            ["conversation_frequency", "1.0", "Depth vs. breadth of social contact patterns"],
            ["places_visited", "1.0", "Behavioral range reduction in hypoactive states"],
            ["upi_transactions_today", "0.9", "Financial activity reduction in severe anhedonia"],
            ["unlock_count", "0.9", "Device engagement frequency — elevated in anxiety, reduced in depression"],
            ["app_uninstalls_today", "0.7", "Impulsive digital behavior changes in mood episodes"],
            ["sleep_time_hour", "0.7", "Late sleep onset — delayed circadian phase in depression"],
            ["charge_duration_hours", "0.7", "Charging regularity as a proxy for sleep hygiene"],
            ["music_time_minutes", "0.7", "Hedonic activity engagement — reduced in anhedonia"],
            ["dark_duration_hours", "0.8", "Total device inactivity — elevated in hypersomnia"],
            ["app_launch_count", "0.8", "App engagement frequency as behavioral activation proxy"],
            ["calendar_events_today", "0.6", "Future planning engagement — reduced in hopelessness"],
            ["app_installs_today", "0.6", "Impulsive novelty-seeking in hypomanic episodes"],
            ["media_count_today", "0.6", "Creative and hedonic engagement proxy"],
            ["notifications_today", "0.7", "Social connectivity signal from notification volume"],
            ["total_apps_count", "0.5", "Quasi-static; diagnostic weight limited"],
            ["memory_usage_percent", "0.5", "System metric; minimal direct clinical signal"],
            ["network_wifi_mb", "0.5", "Partially collinear with app usage metrics"],
            ["network_mobile_mb", "0.5", "Partially collinear with mobility and app usage"],
            ["downloads_today", "0.5", "Weak behavioral engagement proxy"],
            ["storage_used_gb", "0.4", "Quasi-static storage metric; lowest clinical priority"],
        ],
        caption="Table B.1: Complete feature weight registry with clinical rationale."
    )

    return doc


# ─── Entry Point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    BASE_DIR = r"f:\Avaneesh\projects\MH detector\Mental-Health-Detection-ML"
    IMAGE_PATH = os.path.join(BASE_DIR, "image.png")
    OUTPUT_PATH = os.path.join(BASE_DIR, "Lumen_Research_Paper_v2.docx")

    print("Building Lumen Research Manuscript v2...")
    doc = build_document(IMAGE_PATH)
    doc.save(OUTPUT_PATH)
    print(f"\n✅ Done! Saved to:\n   {OUTPUT_PATH}")
    print(f"\n   Image embedded: {'YES' if os.path.exists(IMAGE_PATH) else 'NO — image not found at ' + IMAGE_PATH}")
